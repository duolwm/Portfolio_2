import os
import json
import re
from dataclasses import dataclass
from datetime import datetime
from typing import List, Dict, Tuple

import requests
import streamlit as st
import chromadb
from chromadb.utils import embedding_functions

# ---- Config ----
DB_DIR = "./chroma_db"
BM25_DIR = "./bm25_store"

COL_ESG_REPORTS = "esg_reports"
COL_KNOWLEDGE = "knowledge_bank"
COL_STYLE_REPORT = "style_report"

EMBED_MODEL = "BAAI/bge-small-zh-v1.5"
BM25_STORE = "./bm25_store.jsonl"

OLLAMA_URL = "http://127.0.0.1:11434/api/generate"
LLM_MODEL = "cwchang/llama-3-taiwan-8b-instruct:latest"

def check_ollama_ready() -> tuple[bool, str]:
    try:
        r = requests.get("http://127.0.0.1:11434/api/tags", timeout=5)
        if r.status_code != 200:
            return False, f"Ollama 服務回應狀態碼 {r.status_code}，請確認已啟動 ollama。"
        tags = r.json()
        models = [m.get("name") for m in tags.get("models", []) if isinstance(m, dict)]
        if LLM_MODEL not in models:
            if not any((isinstance(n, str) and LLM_MODEL in n) for n in models):
                return False, f"找不到模型：{LLM_MODEL}\n請先在終端機執行：ollama pull {LLM_MODEL}"
        return True, "OK"
    except requests.exceptions.ConnectionError:
        return False, "無法連線到 Ollama（127.0.0.1:11434）。請確認 Ollama 已啟動。"
    except Exception as e:
        return False, f"Ollama 檢查失敗：{e}"

RERANK_MODEL = "BAAI/bge-reranker-base"

DEFAULT_VEC_K = 12
DEFAULT_BM25_K = 20
DEFAULT_FINAL_K = 8
MAX_CTX_CHARS_PER_CHUNK = 900

TITLE_RE = re.compile(r"^\s*TITLE:\s*(.+)\s*$", re.IGNORECASE | re.MULTILINE)

# ---- BM25 ----
def _simple_zh_tokenize(text: str) -> List[str]:
    text = text.lower()
    text = re.sub(r"[^0-9a-z\u4e00-\u9fff]+", " ", text)
    tokens: List[str] = []
    for part in text.split():
        if re.search(r"[\u4e00-\u9fff]", part):
            tokens.extend(list(part))
        else:
            tokens.append(part)
    return [t for t in tokens if t.strip()]

@dataclass
class BM25Row:
    id: str
    text: str
    meta: Dict

# ---- Vector store ----
@st.cache_resource
def get_collections():
    client = chromadb.PersistentClient(path=DB_DIR)
    embed_fn = embedding_functions.SentenceTransformerEmbeddingFunction(model_name=EMBED_MODEL)
    esg_reports = client.get_or_create_collection(name=COL_ESG_REPORTS, embedding_function=embed_fn)
    knowledge = client.get_or_create_collection(name=COL_KNOWLEDGE, embedding_function=embed_fn)
    style_report = client.get_or_create_collection(name=COL_STYLE_REPORT, embedding_function=embed_fn)
    return esg_reports, knowledge, style_report

def vec_search(col, query: str, k: int) -> List[Tuple[str, str, Dict, float]]:
    res = col.query(query_texts=[query], n_results=k)
    ids = res["ids"][0]
    docs = res["documents"][0]
    metas = res["metadatas"][0]
    dists = res.get("distances", [[None] * len(ids)])[0]
    out = []
    for cid, txt, m, d in zip(ids, docs, metas, dists):
        sim = 0.0 if d is None else (1.0 - float(d))
        out.append((cid, txt, m, sim))
    return out

@st.cache_resource
def load_bm25_rows(collection_name: str):
    rows: List[BM25Row] = []
    path = os.path.join(BM25_DIR, f"{collection_name}.jsonl")
    if not os.path.exists(path):
        return rows
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            obj = json.loads(line)
            rows.append(BM25Row(id=obj["id"], text=obj["text"], meta=obj.get("meta") or {}))
    return rows

@st.cache_resource
def build_bm25_index(collection_name: str):
    from rank_bm25 import BM25Okapi
    rows = load_bm25_rows(collection_name)
    corpus_tokens = [_simple_zh_tokenize(r.text) for r in rows]
    bm25 = BM25Okapi(corpus_tokens) if corpus_tokens else None
    return bm25

def bm25_search(collection_name: str, query: str, k: int) -> List[Tuple[str, str, Dict, float]]:
    rows = load_bm25_rows(collection_name)
    if not rows:
        return []
    bm25 = build_bm25_index(collection_name)
    if bm25 is None:
        return []
    q_tokens = _simple_zh_tokenize(query)
    scores = bm25.get_scores(q_tokens)
    idxs = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:k]
    out = []
    for i in idxs:
        r = rows[i]
        out.append((r.id, r.text, r.meta, float(scores[i])))
    return out

# ---- LLM helpers ----
def ollama_generate(prompt: str, temperature: float = 0.3, max_tokens: int = None) -> str:
    options = {"temperature": temperature}
    if max_tokens:
        options["num_predict"] = max_tokens
        
    r = requests.post(
        OLLAMA_URL,
        json={"model": LLM_MODEL, "prompt": prompt, "stream": False, "options": options},
        timeout=900,
    )
    r.raise_for_status()
    return r.json()["response"]

def hyde_query(topic: str) -> str:
    prompt = f"""你是企業 ESG 與永續發展報告的資深顧問。請根據主題生成一段「假想的標準答案摘要」（200~260字），
用來幫助檢索相關資料。要求：
- 只寫摘要，不要條列，不要引用來源
- 內容要涵蓋 ESG 常見相關構面：例如 環境(E)的管理與排放目標、社會(S)的人權與勞工權益、治理(G)的風險管控等
主題：{topic}
"""
    return ollama_generate(prompt, temperature=0.4).strip()

# ---- Reranker ----
@st.cache_resource
def get_reranker():
    from sentence_transformers import CrossEncoder
    return CrossEncoder(RERANK_MODEL)

def rerank_with_crossencoder(query: str, candidates: List[Tuple[str, str, Dict, float]], top_k: int):
    if not candidates:
        return []
    reranker = get_reranker()
    pairs = [(query, c[1][:1500]) for c in candidates]
    scores = reranker.predict(pairs)
    scored = []
    for c, s in zip(candidates, scores):
        scored.append((c[0], c[1], c[2], float(s)))
    scored.sort(key=lambda x: x[3], reverse=True)
    return scored[:top_k]

def merge_hybrid(vec_hits, bm25_hits, vec_weight: float, bm25_weight: float):
    by_id = {}
    def norm_scores(items):
        if not items:
            return {}
        scores = [it[3] for it in items]
        mn, mx = min(scores), max(scores)
        if mx - mn < 1e-9:
            return {it[0]: 1.0 for it in items}
        return {it[0]: (it[3] - mn) / (mx - mn) for it in items}

    vec_norm = norm_scores(vec_hits)
    bm_norm = norm_scores(bm25_hits)

    for cid, txt, meta, _ in vec_hits:
        by_id[cid] = {"id": cid, "text": txt, "meta": meta, "vec": vec_norm.get(cid, 0.0), "bm25": 0.0}
    for cid, txt, meta, _ in bm25_hits:
        if cid in by_id:
            by_id[cid]["bm25"] = bm_norm.get(cid, 0.0)
        else:
            by_id[cid] = {"id": cid, "text": txt, "meta": meta, "vec": 0.0, "bm25": bm_norm.get(cid, 0.0)}

    merged = []
    for cid, obj in by_id.items():
        score = vec_weight * obj["vec"] + bm25_weight * obj["bm25"]
        merged.append((cid, obj["text"], obj["meta"], float(score)))

    merged.sort(key=lambda x: x[3], reverse=True)
    return merged

# ---- Fact extraction & ChatGPT fact-check prompt ----
def build_fact_prompt(hits, topic: str) -> str:
    ctx_lines = []
    for i, (cid, txt, m, _) in enumerate(hits, start=1):
        txt = (txt[:MAX_CTX_CHARS_PER_CHUNK] + "…") if len(txt) > MAX_CTX_CHARS_PER_CHUNK else txt
        src = f"{m.get('source')}｜{m.get('section_path')}｜part {m.get('part')}"
        ctx_lines.append(f"[{i}] {src}\n{txt}")
    context = "\n\n".join(ctx_lines)

    return f"""你是資料校對員。請只根據【參考內容】抽取「可被驗證的事實」。
目標：讓後續撰寫 ESG 分析報告可以自由發揮，但事實必須準確。

規則：
- 只輸出條列式「事實清單」，不要寫文章、不要下結論。
- 每一條事實末尾必須加上來源引用，例如 [1] 或 [1][3]。
- 不得新增參考內容沒有提到的具體數字、排放量、目標年份、法規或框架要求。
- 若參考內容用語偏經驗/預估（例如『預計將達成』『積極推動』），請保留其預期性表述或加上時間點。

主題：{topic}

【參考內容】
{context}
"""

def build_gpt_factcheck_prompt(facts: str, topic: str) -> str:
    return f"""你是一位「ESG 與永續發展內容事實查核員」，請上網查證下方企業的 ESG 事實清單是否正確，並在必要時修正數字或敘述（必須附可信來源）。

查核規則：
- 優先使用：企業官方 ESG/永續報告書、主管機關公告、國際準則機構（GRI, SASB, TCFD）。
- 每一條都要判定：Correct / PartiallyCorrect / Incorrect
- 若 PartiallyCorrect 或 Incorrect：必須提供修正版（revised），並附可追溯來源
- 不要重寫文章，只要逐條查核。
- sources 請提供可點擊 URL。

輸出格式：請只輸出 JSON（不要加任何多餘文字），結構如下：
{{
  "topic": "...",
  "checked_facts": [
    {{
      "id": 1,
      "original": "...",
      "verdict": "Correct|PartiallyCorrect|Incorrect",
      "revised": "",
      "notes": "",
      "sources": ["https://...", "https://..."]
    }}
  ]
}}

主題：{topic}

【待查核事實清單】
{facts}
"""

def build_generate_prompt_from_checked(checked_json: dict, topic: str, style_snippets: str = "") -> str:
    checked = checked_json.get("checked_facts", [])
    lines = []
    for item in checked:
        fid = item.get("id")
        verdict = (item.get("verdict") or "").strip()
        original = (item.get("original") or "").strip()
        revised = (item.get("revised") or "").strip()
        use = revised if (verdict in ["PartiallyCorrect", "Incorrect"] and revised) else original
        sources = item.get("sources") or []
        src_str = " ".join([f"({s})" for s in sources[:3]])
        lines.append(f"- {use} 〔查核：{verdict}｜id={fid}〕 {src_str}".strip())

    fact_block = "\n".join(lines).strip()
    
    style_instruction = ""
    if style_snippets:
        style_instruction = f"""
# Style Reference (MUST FOLLOW)
請強烈模仿以下【風格樣本】的語氣、段落安排與遣詞用字，將事實清單融入其中，產出符合該風格的文章：
【風格樣本】
{style_snippets}
"""

    return f"""# Role
你是一位資深 ESG 與永續發展顧問，文風「穩健、專業、具洞察力」。你的核心任務是將【已查核事實清單】重組為一篇邏輯嚴密、能幫助決策者與利害關係人了解該議題的專業報告。

# Thinking Process (Internal Logic)
1. **主題類型判定**：首先分析事實屬於「永續治理(G)」、「環境(E)」還是「社會(S)」。
2. **語意自適應標題**：針對特定永續主題設定適當的標題。
3. **數據嚴謹性**：所有具體指標（如減碳量、用電百分比、年分、專案金額）必須 100% 來自清單，不得外推。

# Output Format & Structure
1) TITLE: <20字內，針對主題性質生成，專業且有說服力>
2) Markdown 正文：

## 執行摘要 (Executive Summary)
（140–220字：描述該議題的背景、重要性，以及現階段發展情形。）

## 永續目標與推動機制
（請包含 4–6 個小標題 ###。內容應涵蓋：策略框架、衡量指標指標與達成進度。若有不同年度或項目的數據請整理清晰。）

## 挑戰與風險評估
（條列式描述：推動過程中的挑戰、政策限制，或者清單中提到的證據侷限性與潛在風險。）

## 未來展望與行動建議
（條列式描述：具體可操作的下一步，例如：短中長期目標、增強監管、或需擴大投資的領域。）
{style_instruction}

主題：{topic}

【已查核事實清單】
{fact_block}
"""

def parse_title_and_body(text: str, fallback_title: str):
    m = TITLE_RE.search(text)
    if not m:
        return fallback_title, text.strip()
    title = m.group(1).strip()
    body = TITLE_RE.sub("", text, count=1).strip()
    return (title or fallback_title), body

def try_parse_checked_json(raw: str):
    raw = raw.strip()
    if not raw:
        return None, "你尚未貼上查核結果。"
    try:
        return json.loads(raw), None
    except Exception:
        pass
    m = re.search(r"\{[\s\S]*\}\s*$", raw)
    if m:
        try:
            return json.loads(m.group(0)), None
        except Exception as e:
            return None, f"JSON 解析失敗：{e}"
    return None, "看起來不是 JSON。請確保回覆包含完整 JSON。"

def manuscript_to_docx_bytes(title: str, manuscript: str) -> bytes:
    """將文稿轉成 .docx（二進位）"""
    from docx import Document
    import re
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
        
    prev_empty = False
    for line in (manuscript or "").splitlines():
        line_stripped = line.strip()
        if not line_stripped:
            if not prev_empty:
                doc.add_paragraph("")
                prev_empty = True
            continue
            
        prev_empty = False
        m_heading = re.match(r"^(#{1,6})\s+(.*)", line_stripped)
        if m_heading:
            level = len(m_heading.group(1))
            text = m_heading.group(2)
            text = re.sub(r"(\*\*|__)", "", text)
            doc.add_heading(text, level=level)
            continue

        is_list = False
        if re.match(r"^[-•●*]\s+", line_stripped):
            is_list = True
            line_stripped = re.sub(r"^[-•●*]\s+", "", line_stripped)
            
        p = doc.add_paragraph(style="List Bullet" if is_list else None)
        
        parts = re.split(r'(\*\*.*?\*\*|__.*?__)', line_stripped)
        for part in parts:
            if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')):
                p.add_run(part[2:-2]).bold = True
            elif part:
                p.add_run(part)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---- UI ----
st.set_page_config(page_title="ESG Report RAG", layout="wide")
st.title("🌱 ESG 報告書智慧分析 RAG")

ollama_ok, ollama_msg = check_ollama_ready()
if not ollama_ok:
    st.error(ollama_msg)
    st.stop()

col_esg_reports, col_knowledge, col_style = get_collections()

with st.sidebar:
    st.header("檢索與品質設定")
    use_style = st.checkbox("使用 Style (樣本)", value=True)
    use_knowledge = st.checkbox("使用 Knowledge Bank (法規/框架定錨)", value=True)
    
    use_hyde = st.checkbox("啟用 HyDE (提升召回)", value=True)
    use_bm25 = st.checkbox("啟用 BM25 (關鍵字召回)", value=True)
    if use_bm25 and (not os.path.exists(BM25_STORE)):
        st.warning("找不到 BM25 store。請先執行：python ingest.py --mode all")
    use_rerank = st.checkbox("啟用 Reranker (重排序)", value=True)

    vec_k = st.slider("向量召回 k", 6, 30, DEFAULT_VEC_K)
    bm25_k = st.slider("BM25 召回 k", 10, 60, DEFAULT_BM25_K)
    final_k = st.slider("送入抽 Facts 的段落數", 5, 12, DEFAULT_FINAL_K)

    vec_weight = st.slider("Hybrid 向量權重", 0.0, 1.0, 0.65)
    bm25_weight = st.slider("Hybrid BM25 權重", 0.0, 1.0, 0.35)
    show_context = st.checkbox("顯示檢索段落（稽核用）", value=True)

topic = st.text_area(
    "輸入分析主題/需求（以哪家企業、哪一年度為主，或是問特定議題如碳排）",
    height=120,
    placeholder="例如：整理台積電 2023 年永續報告書中關於水資源管理與回收的具體做法與目標。",
)

if "facts" not in st.session_state:
    st.session_state.facts = ""
if "gpt_prompt" not in st.session_state:
    st.session_state.gpt_prompt = ""
if "checked_raw" not in st.session_state:
    st.session_state.checked_raw = ""
if "checked_json" not in st.session_state:
    st.session_state.checked_json = None
if "hits" not in st.session_state:
    st.session_state.hits = []
if "hyde" not in st.session_state:
    st.session_state.hyde = ""

a, b, c = st.columns([1, 1, 1])

with a:
    if st.button("① 檢索 → 本機抽取 Facts", type="primary", disabled=not topic.strip()):
        q = topic.strip()
        hyde_txt = ""
        if use_hyde:
            with st.spinner("HyDE 生成中…"):
                hyde_txt = hyde_query(q)
            st.session_state.hyde = hyde_txt

        vec_hits = vec_search(col_esg_reports, q, vec_k)
        if use_hyde and hyde_txt:
            vec_hits2 = vec_search(col_esg_reports, hyde_txt, max(6, vec_k // 2))
            by = {cid: (cid, txt, meta, sc) for cid, txt, meta, sc in vec_hits}
            for cid, txt, meta, sc in vec_hits2:
                if cid not in by or sc > by[cid][3]:
                    by[cid] = (cid, txt, meta, sc)
            vec_hits = list(by.values())
            vec_hits.sort(key=lambda x: x[3], reverse=True)
            vec_hits = vec_hits[:vec_k]

        bm_hits = bm25_search(COL_ESG_REPORTS, q, bm25_k) if use_bm25 else []
        merged = merge_hybrid(vec_hits, bm_hits, vec_weight=vec_weight, bm25_weight=bm25_weight) if bm_hits else vec_hits
        
        if use_knowledge:
            k_vec_hits = vec_search(col_knowledge, q, 3)
            k_bm_hits = bm25_search(COL_KNOWLEDGE, q, 5) if use_bm25 else []
            k_merged = merge_hybrid(k_vec_hits, k_bm_hits, 0.5, 0.5) if k_bm_hits else k_vec_hits
            merged.extend(k_merged)

        final_hits = merged[: max(final_k, 12)]
        if use_rerank and final_hits:
            with st.spinner("Reranker 重排序中…"):
                final_hits = rerank_with_crossencoder(q, final_hits, top_k=final_k)
        else:
            final_hits = final_hits[:final_k]

        st.session_state.hits = final_hits

        if show_context:
            st.subheader("🔎 最終段落（送入抽 Facts）")
            for i, (cid, txt, m, sc) in enumerate(final_hits, start=1):
                ttl = f"[{i}] {m.get('source')}｜{m.get('section_path')}｜part {m.get('part')}｜score={sc:.3f}"
                with st.expander(ttl):
                    st.write(txt)

        if use_hyde and st.session_state.hyde:
            st.subheader("🧠 HyDE（供你檢視）")
            st.write(st.session_state.hyde)

        st.subheader("🧾 Stage 1：抽取事實清單")
        fact_prompt = build_fact_prompt(final_hits, q)
        try:
            with st.spinner("本機抽取 Facts 中…"):
                st.session_state.facts = ollama_generate(fact_prompt, temperature=0.2)
        except Exception as e:
            st.error(f"事實抽取失敗：{e}")

with b:
    if st.button("② 產生『貼到 ChatGPT』查核 Prompt（只回 JSON）", disabled=not st.session_state.facts.strip()):
        st.session_state.gpt_prompt = build_gpt_factcheck_prompt(st.session_state.facts, topic.strip())

with c:
    if st.button("③ 解析查核結果 → 生成最終報告文稿", disabled=not topic.strip()):
        checked_json, err = try_parse_checked_json(st.session_state.checked_raw)
        if err:
            st.error(err)
        else:
            st.session_state.checked_json = checked_json
            st.subheader("✍️ Stage 3：用『已查核事實』生成報告")
            
            style_snippets = ""
            if use_style:
                with st.spinner("檢索風格樣本中…"):
                    style_hits = vec_search(col_style, topic.strip() or "ESG", k=2)
                    if style_hits:
                        style_snippets = "\n\n---\n\n".join([txt for _, txt, _, _ in style_hits])
                        
            gen_prompt = build_generate_prompt_from_checked(checked_json, topic.strip(), style_snippets)
            try:
                with st.spinner("本機生成報告中…"):
                    draft = ollama_generate(gen_prompt, temperature=0.35)
            except Exception as e:
                st.error(f"生成失敗：{e}")
                st.stop()

            gen_title, body = parse_title_and_body(draft, fallback_title=topic.strip())

            manuscript = body.strip()
            manuscript_clean = re.sub(r'\n{3,}', '\n\n', manuscript)

            os.makedirs("output", exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = ts + "_esg_report"

            txt_name = base_name + ".txt"
            txt_path = os.path.join("output", txt_name)
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(gen_title.strip() + "\n\n" + manuscript_clean.strip() + "\n")

            st.success(f"✅ 已產出：output/{txt_name}")
            st.download_button("下載文稿 (.txt)", data=(gen_title.strip() + "\n\n" + manuscript_clean).encode("utf-8"), file_name=txt_name, mime="text/plain")

            try:
                docx_bytes = manuscript_to_docx_bytes(gen_title.strip(), manuscript_clean)
                docx_name = base_name + ".docx"
                st.download_button("下載文稿 (.docx)", data=docx_bytes, file_name=docx_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.caption(f"（未安裝 python-docx 或轉換失敗，無法提供 docx 檔案。錯誤：{e}）")

            st.subheader("文稿預覽")
            st.write(gen_title.strip())
            st.write(manuscript_clean)

st.divider()

st.subheader("Facts（本機抽取）")
st.code(st.session_state.facts or "（尚未抽取）", language="markdown")

st.subheader("『貼到 ChatGPT』查核 Prompt（直接全段複製貼上）")
if st.session_state.gpt_prompt:
    st.code(st.session_state.gpt_prompt, language="text")
else:
    st.info("按 ② 產生查核 Prompt。")

st.subheader("貼回 ChatGPT 的查核 JSON")
st.session_state.checked_raw = st.text_area(
    "把 ChatGPT 回覆的 JSON 原樣貼在這裡（sources 請包含 URL）",
    value=st.session_state.checked_raw,
    height=260,
    placeholder='{\n  "topic": "...",\n  "checked_facts": [ ... ]\n}\n',
)

if st.session_state.checked_json:
    st.subheader("已解析的查核結果（預覽）")
    st.json(st.session_state.checked_json)
